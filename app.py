"""
Job Finder AI — Streamlit web app.

Run with:  streamlit run app.py
"""

import os
import re
import json
import tempfile
from io import BytesIO
from datetime import datetime

import streamlit as st
import pandas as pd

from scraper import (
    search_jobs, search_jobs_multi, fetch_job_page_texts,
    SUPPORTED_SITES, SITE_LABELS,
)
from cv_analyzer import (
    read_cv, score_cv_against_jobs, extract_skills,
    ai_suggestions, ai_cover_letter, ai_tailored_cv, ai_interview_prep,
)
from excel_exporter import export_to_excel
from locations import (
    ALL_JOB_TITLES, JOB_TITLES, COUNTRY_LIST,
    get_cities, build_location_string,
    ALL_PLATFORMS, PLATFORM_INFO, get_recommended_sites,
)

# ──────────────────────────────────────────────────────────────────────────────
# Page config
# ──────────────────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Job Finder AI",
    page_icon="💼",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ──────────────────────────────────────────────────────────────────────────────
# Helpers  (defined before first use)
# ──────────────────────────────────────────────────────────────────────────────

def _safe_filename(text: str) -> str:
    """Strip characters that are illegal in file names."""
    return re.sub(r'[\\/*?:"<>|]', "_", text).replace(" ", "_")


# ── Local persistence (no database — plain JSON next to the app) ──────────────
_TRACKER_PATH        = "job_tracker.json"
_SEEN_PATH           = "seen_jobs.json"
_SAVED_SEARCHES_PATH = "saved_searches.json"

TRACK_STATUSES = ["Not tracked", "Saved", "Applied", "Interview", "Offer", "Rejected"]


def _load_json(path: str, default):
    try:
        with open(path, encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return default


def _save_json(path: str, data) -> None:
    try:
        with open(path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=1)
    except Exception as e:
        st.warning(f"Could not save {path}: {e}")


def _text_to_docx_bytes(text: str) -> bytes:
    """Render plain CV text (UPPERCASE headings, '- ' bullets) into a .docx."""
    from docx import Document
    doc = Document()
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.isupper() and 2 < len(stripped) < 60:
            doc.add_heading(stripped.title(), level=2)
        elif stripped.startswith(("- ", "• ")):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(stripped)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _csv_safe(df: pd.DataFrame) -> pd.DataFrame:
    """Formula-injection guard for CSV export: Excel executes cells starting
    with = + - @ when the file is opened. Scraped titles/descriptions are
    attacker-controlled, so prefix such cells with ' (rendered as text)."""
    df = df.copy()
    # include "string": pandas ≥3 infers str columns as StringDtype, not object
    for col in df.select_dtypes(include=["object", "string"]).columns:
        df[col] = df[col].map(
            lambda v: "'" + v if isinstance(v, str) and v[:1] in "=+-@" else v
        )
    return df


@st.cache_data(ttl=900, show_spinner=False)
def _cached_search(keywords: tuple, location: str, max_results: int, hours_old: int,
                   sites: tuple, remote_only: bool, fetch_desc: bool,
                   _desc_workers: int, country: str) -> pd.DataFrame:
    """Repeating an identical search within 15 min returns instantly instead
    of re-scraping every board. _desc_workers is underscore-prefixed so it is
    excluded from the cache key — it changes fetch speed, not results, and
    moving the slider must not bust the cache."""
    return search_jobs_multi(
        keywords=list(keywords),
        location=location,
        max_results=max_results,
        hours_old=hours_old,
        sites=list(sites),
        remote_only=remote_only,
        fetch_full_descriptions=fetch_desc,
        desc_workers=_desc_workers,
        country=country,
    )


def _attach_job_skills(df: pd.DataFrame) -> pd.DataFrame:
    """
    Add a job_skills column (sorted list of detected skills per description).
    Computed once at search time so reruns don't re-scan every description.
    """
    df = df.copy()
    df["job_skills"] = [
        sorted(extract_skills(str(d or ""))) for d in df["description"].fillna("")
    ]
    return df


# Fallback search URLs (from _ensure_job_urls) land on a results page, not a
# job — fetching them would score skills from unrelated postings.
_SEARCH_URL_MARKERS = (
    "/jobs/search", "jobs?q=", "jobs-search?k=", "jobs?keywords=",
    "/?text=", "/jsearch/", "google.com/search",
)


def _is_direct_job_url(url: str) -> bool:
    u = str(url or "").lower()
    return u.startswith(("http://", "https://")) and not any(m in u for m in _SEARCH_URL_MARKERS)


def _fill_missing_job_skills(df: pd.DataFrame) -> pd.DataFrame:
    """
    Boards like Bayt/Reed/Rozee/Indeed return cards with no description, so
    their job_skills is empty and the ATS score would always be 0. For those
    rows, fetch the job page once, keep ONLY the extracted skill keywords,
    and discard the page text — it is never stored or shown in the UI.
    Called only when a CV is loaded (the fetches exist purely to score it).
    """
    if df.empty or "job_skills" not in df.columns:
        return df

    descs = df["description"].fillna("").astype(str).tolist()
    urls  = df["job_url"].fillna("").astype(str).tolist()
    need  = [
        i for i, (d, sk, u) in enumerate(zip(descs, df["job_skills"], urls))
        if not d.strip() and not sk and _is_direct_job_url(u)
    ]
    if not need:
        return df

    texts = fetch_job_page_texts([urls[i] for i in need])

    df = df.copy()
    skills_col = list(df["job_skills"])
    for i, text in zip(need, texts):
        if text:
            skills_col[i] = sorted(extract_skills(text))
    df["job_skills"] = skills_col          # small keyword lists — no page text
    return df


def _compute_scores(cv_text: str, df: pd.DataFrame, cv_skills: set = None,
                    api_key: str = None, provider: str = None) -> pd.DataFrame:
    """
    Add match_score, ats_score, matched_skills, missing_skills columns.
    Operates on the FULL dataframe before any UI filter is applied — this
    keeps scores aligned with the integer row index used later.

    ats_score = % of the skills detected in the job description that also
    appear in the CV — the keyword-coverage signal an ATS keyword screen uses,
    as opposed to match_score's free-text similarity.
    """
    descriptions = df["description"].fillna("").tolist()
    sim_scores   = score_cv_against_jobs(cv_text, descriptions,
                                         api_key=api_key, provider=provider)
    if cv_skills is None:
        cv_skills = extract_skills(cv_text)

    if "job_skills" in df.columns:
        job_skills = [set(sk or []) for sk in df["job_skills"]]
    else:
        job_skills = [extract_skills(str(d)) for d in descriptions]

    matched = [sk & cv_skills for sk in job_skills]

    df = df.copy()
    df["match_score"]    = sim_scores
    df["matched_skills"] = [sorted(m) for m in matched]
    df["missing_skills"] = [sorted(sk - cv_skills) for sk in job_skills]
    df["ats_score"]      = [
        round(len(m) / len(sk) * 100, 1) if sk else 0.0
        for m, sk in zip(matched, job_skills)
    ]
    return df


# ── Experience-requirement extraction ────────────────────────────────────────
# Parsed once at search time; jobs that don't state a requirement get NaN and
# always pass the experience filter (same convention as the date filter).

_EXP_RANGE_RE  = re.compile(
    r"(\d{1,2})\s*(?:-|–|—|to)\s*(\d{1,2})\s*\+?\s*(?:years?|yrs?)\b", re.IGNORECASE)
_EXP_PLUS_RE   = re.compile(r"(\d{1,2})\s*\+\s*(?:years?|yrs?)\b", re.IGNORECASE)
_EXP_MIN_RE    = re.compile(
    r"(?:at least|minimum|min\.?)\s+(\d{1,2})\s*(?:years?|yrs?)\b", re.IGNORECASE)
_EXP_SINGLE_RE = re.compile(
    r"(\d{1,2})\s*(?:years?|yrs?)['’]?\s*(?:of\s+)?(?:[\w/+-]+\s+){0,3}experience",
    re.IGNORECASE)

_EXP_OPEN_MAX = 99   # stand-in upper bound for open-ended "5+ years"


def _extract_exp_range(text: str):
    """Return (min_years, max_years) required by a job posting, or None.
    "5+ years" / "minimum 5 years" / "5 years of experience" are treated as
    open-ended minimums."""
    if not text:
        return None
    m = _EXP_RANGE_RE.search(text)
    if m:
        lo, hi = int(m.group(1)), int(m.group(2))
        return (lo, hi) if lo <= hi else (hi, lo)
    m = _EXP_PLUS_RE.search(text) or _EXP_MIN_RE.search(text) or _EXP_SINGLE_RE.search(text)
    if m:
        return (int(m.group(1)), _EXP_OPEN_MAX)
    return None


def _attach_experience(df: pd.DataFrame) -> pd.DataFrame:
    """Add exp_min / exp_max columns parsed from title + job_level + description
    (job_level covers Naukri, which puts "2-7 Yrs" there)."""
    df = df.copy()
    blobs = (
        df["title"].fillna("").astype(str) + " · "
        + df["job_level"].fillna("").astype(str) + " · "
        + df["description"].fillna("").astype(str)
    )
    ranges = [_extract_exp_range(b) for b in blobs]
    df["exp_min"] = [r[0] if r else None for r in ranges]
    df["exp_max"] = [r[1] if r else None for r in ranges]
    return df


# ──────────────────────────────────────────────────────────────────────────────
# Session state
# ──────────────────────────────────────────────────────────────────────────────
for _k, _v in {
    "jobs_df":     None,
    "cv_text":     None,
    "cv_skills":   None,
    "cv_file_id":  None,
    "excel_bytes": None,
}.items():
    if _k not in st.session_state:
        st.session_state[_k] = _v

# ──────────────────────────────────────────────────────────────────────────────
# Sidebar — search controls
# ──────────────────────────────────────────────────────────────────────────────
with st.sidebar:
    st.title("💼 Job Finder AI")
    st.caption("Multi-platform job search — LinkedIn, Indeed, Remotive, Jobicy, The Muse, Arbeitnow, Naukri, Seek, JobStreet, Bayt, Reed, Rozee & more.")
    st.divider()

    # ── Job Titles (multiple) ─────────────────────────────────────────────────
    st.subheader("🔍 Job Titles")

    selected_from_list = st.multiselect(
        "Pick from predefined list",
        options=ALL_JOB_TITLES,
        placeholder="Type to search and select…",
        label_visibility="collapsed",
        help="Search by typing — select as many titles as you like.",
        key="kw_select",
    )

    custom_raw = st.text_area(
        "Add custom titles (one per line)",
        placeholder="Senior Python Developer\nML Engineer\nData Scientist",
        height=72,
        help="Enter job titles not in the list above, one per line.",
        key="custom_titles",
    )
    custom_titles = [t.strip() for t in custom_raw.splitlines() if t.strip()]

    # Merge list + custom, deduplicate, preserve order
    keywords = list(dict.fromkeys(selected_from_list + custom_titles))

    if keywords:
        st.caption(
            f"**{len(keywords)} title(s):** "
            + "  ·  ".join(f"`{k}`" for k in keywords)
        )
    else:
        st.caption("No job title selected yet — pick from list or type above.")

    st.divider()

    # ── Location — country → city ─────────────────────────────────────────────
    st.subheader("📍 Location")

    country = st.selectbox("Country / Region", COUNTRY_LIST, key="country_sel")
    cities  = get_cities(country)

    if country == "Remote (Worldwide)":
        city = st.selectbox("Region", cities, key="city_sel")
    elif cities:
        # "Entire country" first and default — a silently pre-selected first
        # city made users think they were searching the whole country.
        city_options = ["Entire country (no city filter)"] + cities
        city = st.selectbox("City", city_options, key="city_sel")
        if city == "Entire country (no city filter)":
            city = ""
    else:
        city = ""

    location = build_location_string(country, city)
    st.caption(f"Search location: **{location}**")

    st.divider()

    # ── Search options ────────────────────────────────────────────────────────
    st.subheader("⚙️ Search Options")

    recommended = get_recommended_sites(country)
    platform_labels = {k: v["name"] for k, v in PLATFORM_INFO.items()}

    sites = st.multiselect(
        "Job Boards",
        options=ALL_PLATFORMS,
        default=recommended,
        format_func=lambda s: platform_labels.get(s, s),
        key=f"sites_{country}",  # re-renders when country changes
        help="Boards are pre-selected based on the chosen country.",
    )

    col1, col2 = st.columns(2)
    max_results = col1.number_input("Max Jobs", min_value=5, max_value=100, value=30, step=5,
                                    key="max_results")
    _HOURS_OPTS = [1, 2, 4, 8, 24, 48, 72, 168, 336, 720]
    hours_old   = col2.selectbox(
        "Posted Within",
        _HOURS_OPTS,
        index=7,
        format_func=lambda h: f"{h}h" if h < 24 else f"{h // 24}d",
        key="hours_old",
    )

    # ── Experience range (optional search parameter) ──────────────────────────
    use_exp = st.checkbox(
        "Experience range (years)  *(optional)*",
        value=False,
        help="Return only jobs whose required experience falls in this range. "
             "Jobs that don't state a requirement are kept too.",
        key="use_exp",
    )
    exp_min_in = exp_max_in = None
    if use_exp:
        ec1, ec2 = st.columns(2)
        exp_min_in = ec1.number_input("Min years", min_value=0, max_value=50, value=0,
                                      key="exp_min_in")
        exp_max_in = ec2.number_input("Max years", min_value=0, max_value=50, value=5,
                                      key="exp_max_in")
        if exp_min_in > exp_max_in:
            exp_min_in, exp_max_in = exp_max_in, exp_min_in

    # Reset "Remote only" whenever the country changes — with an explicit
    # widget key the old value silently persists otherwise (e.g. staying
    # checked after switching from Remote (Worldwide) to Netherlands, which
    # turns an entire-country search into country-wide remote listings only).
    if st.session_state.get("_last_country") != country:
        st.session_state["remote_only"] = (country == "Remote (Worldwide)")
        st.session_state["_last_country"] = country

    remote_only = st.checkbox("Remote only", value=(country == "Remote (Worldwide)"),
                              key="remote_only")
    fetch_desc  = st.checkbox(
        "Fetch full descriptions (LinkedIn)",
        value=True,
        help="Needed for CV match scoring. Makes search ~2× slower.",
    )
    desc_workers = st.slider(
        "Parallel description fetches",
        min_value=1, max_value=10, value=5,
        help="Higher = faster but more likely to be rate-limited.",
    )

    st.divider()

    # ── CV upload ─────────────────────────────────────────────────────────────
    st.subheader("📄 CV Upload  *(optional)*")
    cv_file = st.file_uploader(
        "Upload your CV",
        type=["docx", "pdf"],
        help="Get a match score for every job.",
    )

    st.divider()

    # ── AI Suggestions ────────────────────────────────────────────────────────
    st.subheader("🤖 AI Suggestions  *(optional)*")

    ai_provider = st.selectbox(
        "Provider",
        options=["Claude (Anthropic)", "ChatGPT (OpenAI)", "Gemini (Google)"],
        help="Pick which AI gives you CV improvement tips.",
    )

    _PROVIDER_META = {
        "Claude (Anthropic)": ("claude", "sk-ant-…", "ANTHROPIC_API_KEY"),
        "ChatGPT (OpenAI)":   ("openai", "sk-…",     "OPENAI_API_KEY"),
        "Gemini (Google)":    ("gemini", "AIza…",    "GOOGLE_API_KEY"),
    }
    provider_key, _ph, _env_var = _PROVIDER_META[ai_provider]

    api_key = st.text_input(
        "API Key",
        type="password",
        placeholder=_ph,
        help=f"Leave blank — AI suggestions are fully optional. "
             f"Also picked up from the {_env_var} environment variable.",
    )
    # Env-var keys count too — previously the suggestion button never
    # appeared unless the key was typed into the sidebar.
    has_ai_key = bool(api_key or os.getenv(_env_var))

    st.divider()

    # ── Saved searches ────────────────────────────────────────────────────────
    with st.expander("💾 Saved Searches"):
        _saved = _load_json(_SAVED_SEARCHES_PATH, {})

        if _saved:
            _sel = st.selectbox("Load a saved search", sorted(_saved.keys()))
            lc1, lc2 = st.columns(2)
            if lc1.button("Load", use_container_width=True):
                cfg = _saved.get(_sel, {})
                st.session_state["kw_select"] = [
                    t for t in cfg.get("titles", []) if t in ALL_JOB_TITLES
                ]
                st.session_state["custom_titles"] = cfg.get("custom", "")
                _ctry = cfg.get("country", COUNTRY_LIST[0])
                if _ctry in COUNTRY_LIST:
                    st.session_state["country_sel"] = _ctry
                    # city must be valid for the loaded country or Streamlit errors
                    _cs = get_cities(_ctry)
                    _opts = _cs if _ctry == "Remote (Worldwide)" else (
                        ["Entire country (no city filter)"] + _cs if _cs else []
                    )
                    if cfg.get("city") in _opts:
                        st.session_state["city_sel"] = cfg["city"]
                    else:
                        st.session_state.pop("city_sel", None)
                    _sites = [s for s in cfg.get("sites", []) if s in ALL_PLATFORMS]
                    if _sites:
                        st.session_state[f"sites_{_ctry}"] = _sites
                if isinstance(cfg.get("max_results"), int):
                    st.session_state["max_results"] = min(max(cfg["max_results"], 5), 100)
                if cfg.get("hours_old") in _HOURS_OPTS:
                    st.session_state["hours_old"] = cfg["hours_old"]
                st.session_state["remote_only"] = bool(cfg.get("remote_only", False))
                # keep the country-change reset from clobbering the loaded value
                st.session_state["_last_country"] = cfg.get("country", "")
                st.session_state["use_exp"]     = bool(cfg.get("use_exp", False))
                if cfg.get("use_exp"):
                    st.session_state["exp_min_in"] = int(cfg.get("exp_min", 0))
                    st.session_state["exp_max_in"] = int(cfg.get("exp_max", 5))
                st.rerun()
            if lc2.button("Delete", use_container_width=True):
                _saved.pop(_sel, None)
                _save_json(_SAVED_SEARCHES_PATH, _saved)
                st.rerun()

        _save_name = st.text_input("Save current settings as", placeholder="e.g. Python remote")
        if st.button("Save current search", use_container_width=True):
            if not _save_name.strip():
                st.warning("Give the search a name first.")
            else:
                _saved[_save_name.strip()] = {
                    "titles":      selected_from_list,
                    "custom":      custom_raw,
                    "country":     country,
                    "city":        st.session_state.get("city_sel", ""),
                    "sites":       sites,
                    "max_results": int(max_results),
                    "hours_old":   hours_old,
                    "remote_only": remote_only,
                    "use_exp":     use_exp,
                    "exp_min":     int(exp_min_in or 0),
                    "exp_max":     int(exp_max_in or 0),
                }
                _save_json(_SAVED_SEARCHES_PATH, _saved)
                st.success(f"Saved “{_save_name.strip()}”")

    search_btn = st.button("Search Jobs  🚀", type="primary", use_container_width=True)

# ──────────────────────────────────────────────────────────────────────────────
# CV upload — process immediately so scores are ready after search
# ──────────────────────────────────────────────────────────────────────────────
if cv_file:
    # Only re-parse when the user uploads a different file; skip on every rerun.
    if cv_file.file_id != st.session_state.get("cv_file_id"):
        suffix = "." + cv_file.name.rsplit(".", 1)[-1]
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(cv_file.read())
            tmp_path = tmp.name
        try:
            cv_text = read_cv(tmp_path)
        finally:
            os.unlink(tmp_path)  # don't leave CV copies in %TEMP% on a parse error

        if cv_text:
            st.session_state.cv_text   = cv_text
            # Extract once here — scanning ~180 skill regexes over the CV on
            # every rerun (each widget click) was measurable latency.
            st.session_state.cv_skills  = extract_skills(cv_text)
            st.session_state.cv_file_id = cv_file.file_id
            st.sidebar.success(f"CV loaded ✓  ({len(cv_text):,} chars)")
            # Re-score immediately if jobs already loaded
            if st.session_state.jobs_df is not None:
                with st.spinner("Computing ATS & match scores …"):
                    st.session_state.jobs_df = _compute_scores(
                        cv_text,
                        _fill_missing_job_skills(st.session_state.jobs_df),
                        cv_skills=st.session_state.cv_skills,
                        api_key=api_key, provider=provider_key,
                    )
        else:
            st.sidebar.error("Could not read CV — check the file format.")
    else:
        st.sidebar.success(f"CV loaded ✓  ({len(st.session_state.cv_text or ''):,} chars)")
elif st.session_state.cv_file_id is not None:
    # CV removed via the uploader's ✕ — drop the stale text and scores
    st.session_state.cv_text    = None
    st.session_state.cv_skills  = None
    st.session_state.cv_file_id = None
    if st.session_state.jobs_df is not None:
        st.session_state.jobs_df = st.session_state.jobs_df.drop(
            columns=["match_score", "ats_score", "matched_skills", "missing_skills"],
            errors="ignore",
        )

# ──────────────────────────────────────────────────────────────────────────────
# Search
# ──────────────────────────────────────────────────────────────────────────────
if search_btn:
    if not keywords:
        st.error("Please select or type at least one job title.")
    elif not location.strip():
        st.error("Please select a location.")
    elif not sites:
        st.error("Select at least one job board.")
    else:
        status = st.empty()
        kw_display = ", ".join(f"**{k}**" for k in keywords[:3])
        if len(keywords) > 3:
            kw_display += f" *+{len(keywords) - 3} more*"
        status.info(
            f"Searching **{', '.join(s.capitalize() for s in sites)}** "
            f"for {kw_display} in **{location}** …  "
            f"({len(keywords)} title(s) × {len(sites)} board(s) in parallel)"
        )

        with st.spinner("Fetching jobs …"):
            df = _cached_search(
                keywords=tuple(keywords),
                location=location,
                max_results=max_results,
                hours_old=hours_old,
                sites=tuple(sites),
                remote_only=remote_only,
                fetch_desc=fetch_desc,
                _desc_workers=desc_workers,
                country=country,
            )

        status.empty()

        # Experience-range search parameter: parse each posting's stated
        # requirement and keep jobs in the requested range (unstated pass).
        _emptied_by_exp = False
        if df is not None and not df.empty:
            df = _attach_experience(df)
            if use_exp:
                _no_req  = df["exp_min"].isna()
                _overlap = (df["exp_min"] <= exp_max_in) & (df["exp_max"] >= exp_min_in)
                df = df[_no_req | _overlap].reset_index(drop=True)
                _emptied_by_exp = df.empty

        if _emptied_by_exp:
            st.warning(
                f"**Jobs were found, but none match the {exp_min_in}–{exp_max_in} "
                "years experience range.** Widen the range or uncheck the "
                "experience option in the sidebar."
            )
            st.session_state.jobs_df = None
        elif df is None or df.empty:
            st.warning(
                "**No jobs found.** Common reasons:\n\n"
                "- The site may be temporarily blocking requests — wait 30 s and retry\n"
                "- Try **broader keywords** (e.g. `developer` not `senior python developer`)\n"
                "- Try a **larger location** (entire country, not just a city)\n"
                "- Increase **Posted Within** to 7d or 30d\n"
                "- Try **LinkedIn only** — it is the most reliable source\n"
                "- Country-specific boards (Naukri, Seek, JobStreet, Bayt, Reed, Rozee) "
                "may block automated requests; LinkedIn is the safest fallback"
            )
            st.session_state.jobs_df = None
        else:
            df = _attach_job_skills(df)

            # Mark jobs never seen in any previous search (🆕 badge + filter)
            _seen = _load_json(_SEEN_PATH, {})
            _urls = df["job_url"].fillna("").astype(str).tolist()
            df["is_new"] = [u not in _seen for u in _urls]
            _now = datetime.now().strftime("%Y-%m-%d %H:%M")
            for _u in _urls:
                _seen.setdefault(_u, _now)
            if len(_seen) > 5000:   # cap file growth — keep the newest entries
                _seen = dict(sorted(_seen.items(), key=lambda kv: kv[1])[-5000:])
            _save_json(_SEEN_PATH, _seen)

            # Attach match scores as columns (avoids index-alignment bugs).
            # Boards without card descriptions get their skills pulled from
            # the job page on the fly (text discarded) so ATS isn't 0.
            if st.session_state.cv_text:
                with st.spinner("Computing ATS & match scores …"):
                    df = _fill_missing_job_skills(df)
                    df = _compute_scores(
                        st.session_state.cv_text, df,
                        cv_skills=st.session_state.cv_skills,
                        api_key=api_key, provider=provider_key,
                    )

            # Clear per-row results from the previous search — they're keyed by
            # row index and would show up under unrelated jobs otherwise.
            _ROW_KEY_PREFIXES = ("ai_res_", "cl_res_", "ip_res_", "cv_docx_", "trk_")
            for _k in [k for k in st.session_state
                       if str(k).startswith(_ROW_KEY_PREFIXES)]:
                del st.session_state[_k]

            st.session_state.jobs_df    = df
            st.session_state.excel_bytes = None
            boards_found = ", ".join(df["site"].unique())
            kw_summary   = ", ".join(f'"{k}"' for k in keywords[:3])
            if len(keywords) > 3:
                kw_summary += f" +{len(keywords) - 3} more"
            _n_new = int(df["is_new"].sum()) if "is_new" in df.columns else 0
            st.success(
                f"Found **{len(df)} jobs** for {kw_summary} "
                f"from {boards_found}"
                + (f" — **{_n_new} new** since your last searches." if _n_new else ".")
            )

            # Board health — a board at 0 is blocked, keyless, or has no matches
            _counts = df["site"].value_counts()
            _parts = []
            for _s in sites:
                _label = SITE_LABELS.get(_s, _s)
                _n = int(_counts.get(_label, 0))
                _parts.append(f"{_label} **{_n}** ✓" if _n else f"{_label} **0** ⚠️")
            st.caption(
                "Board results: " + " · ".join(_parts)
                + "  — ⚠️ = blocked, missing API key, or no matches on that board."
            )

# ──────────────────────────────────────────────────────────────────────────────
# Main content
# ──────────────────────────────────────────────────────────────────────────────
df = st.session_state.jobs_df
has_scores = df is not None and "match_score" in df.columns

if df is not None and not df.empty:

    tab_list, tab_analysis, tab_tracker, tab_download = st.tabs(
        ["📋 Job List", "📊 CV Analysis", "📌 Tracker", "📥 Download"]
    )

    # Tracker state shared by the Job List cards and the Tracker tab
    tracker = _load_json(_TRACKER_PATH, {})

    # ── Tab 1 — Job List ─────────────────────────────────────────────────────
    with tab_list:
        # Filter row
        fc1, fc2, fc3, fc4, fc5 = st.columns([2, 2, 2, 2, 2])

        kw_col_exists = "searched_keyword" in df.columns
        kw_opts  = sorted(df["searched_keyword"].dropna().unique().tolist()) if kw_col_exists else []
        filter_kw = fc1.multiselect("Job Title", kw_opts) if kw_opts else []

        site_opts   = df["site"].unique().tolist() if "site" in df.columns else []
        filter_site = fc2.multiselect("Site", site_opts)

        filter_remote = fc3.selectbox("Remote", ["All", "Remote Only", "On-site Only"])

        type_opts   = [t for t in df["job_type"].unique() if t] if "job_type" in df.columns else []
        filter_type = fc4.multiselect("Job Type", type_opts)

        min_score = fc5.slider("Min Match %", 0, 100, 0) if has_scores else 0

        # Sort + new-only row
        sc1, sc2, _sp = st.columns([2, 2, 6])
        _sort_opts = ["Default"]
        if has_scores:
            _sort_opts += ["Match %", "ATS %"]
        _sort_opts += ["Newest", "Salary"]
        sort_by = sc1.selectbox("Sort by", _sort_opts)
        only_new = (
            sc2.checkbox("🆕 New only", help="Jobs never returned by a previous search.")
            if "is_new" in df.columns else False
        )

        # Build a boolean mask so we never copy the whole DataFrame per filter change
        mask = pd.Series(True, index=df.index)
        if filter_kw and kw_col_exists:
            mask &= df["searched_keyword"].isin(filter_kw)
        if filter_site:
            mask &= df["site"].isin(filter_site)
        if filter_remote == "Remote Only" and "is_remote" in df.columns:
            mask &= df["is_remote"] == True
        elif filter_remote == "On-site Only" and "is_remote" in df.columns:
            mask &= df["is_remote"] != True
        if filter_type and "job_type" in df.columns:
            mask &= df["job_type"].isin(filter_type)
        if has_scores and min_score > 0:
            mask &= df["match_score"] >= min_score
        if only_new:
            mask &= df["is_new"] == True
        view = df[mask]

        if sort_by == "Match %":
            view = view.sort_values("match_score", ascending=False)
        elif sort_by == "ATS %" and "ats_score" in view.columns:
            view = view.sort_values("ats_score", ascending=False)
        elif sort_by == "Newest":
            # ISO date strings sort lexicographically; empty dates go last
            view = view.sort_values("date_posted", ascending=False, na_position="last")
        elif sort_by == "Salary" and "max_amount" in view.columns:
            _sal_key = pd.to_numeric(view["max_amount"], errors="coerce").fillna(
                pd.to_numeric(view["min_amount"], errors="coerce")
            )
            view = view.assign(_sal=_sal_key).sort_values(
                "_sal", ascending=False, na_position="last"
            ).drop(columns="_sal")

        # Paginate — rendering 100 expanders on every rerun (each widget click
        # re-runs the whole script) was the main source of UI lag.
        PAGE_SIZE = 15
        if len(view) > PAGE_SIZE:
            n_pages = -(-len(view) // PAGE_SIZE)
            pg_col, cap_col = st.columns([1, 5])
            page_no = pg_col.selectbox(
                "Page", list(range(1, n_pages + 1)),
                format_func=lambda p: f"Page {p} / {n_pages}",
                label_visibility="collapsed",
            )
            lo_i = (page_no - 1) * PAGE_SIZE
            cap_col.caption(
                f"Showing **{lo_i + 1}–{min(lo_i + PAGE_SIZE, len(view))}** "
                f"of **{len(view)}** filtered jobs ({len(df)} total)"
            )
            view = view.iloc[lo_i:lo_i + PAGE_SIZE]
        else:
            st.caption(f"Showing **{len(view)}** of **{len(df)}** jobs")

        # CV skills cached at upload time — not re-extracted per rerun
        if st.session_state.cv_text and st.session_state.cv_skills is None:
            st.session_state.cv_skills = extract_skills(st.session_state.cv_text)
        cv_skills = st.session_state.cv_skills or set()

        # Job cards
        for _, row in view.iterrows():
            pct = row.get("match_score", 0) or 0
            badge = "🟢" if pct >= 65 else ("🟡" if pct >= 35 else ("🔴" if pct > 0 else "⚪"))

            searched_kw = row.get("searched_keyword", "")
            ats = row.get("ats_score")
            has_ats = ats is not None and not pd.isna(ats)

            label = f"{badge} **{row.get('title', 'N/A')}** — {row.get('company', 'N/A')}"
            if searched_kw:
                label += f"  · `{searched_kw}`"
            if pct:
                label += f"  `{pct:.0f}% match`"
            if has_ats:
                label += f"  `ATS {ats:.0f}%`"
            if row.get("is_new"):
                label += "  🆕"
            _trk_status = tracker.get(str(row.get("job_url") or ""), {}).get("status")
            if _trk_status:
                label += f"  `📌 {_trk_status}`"

            # Stated experience requirement (parsed at search time)
            _emin, _emax = row.get("exp_min"), row.get("exp_max")
            if _emin is not None and not pd.isna(_emin):
                exp_str = (
                    f"{_emin:.0f}+ yrs" if _emax == 99
                    else f"{_emin:.0f} yrs" if _emin == _emax
                    else f"{_emin:.0f}–{_emax:.0f} yrs"
                )
            else:
                exp_str = "—"

            with st.expander(label, expanded=False):
                m1, m2, m3 = st.columns(3)
                m1.markdown(f"📍 **Location:** {row.get('location', '—')}")
                m1.markdown(f"🏢 **Site:** {row.get('site', '—')}")
                m1.markdown(f"🔎 **Searched:** {searched_kw or '—'}")
                m2.markdown(f"📅 **Posted:** {str(row.get('date_posted') or '—')[:10]}")
                m2.markdown(f"💼 **Type:** {row.get('job_type', '—') or '—'}")
                m2.markdown(f"⏳ **Experience:** {exp_str}")
                m3.markdown(f"📊 **Level:** {row.get('job_level', '—') or '—'}")
                m3.markdown(f"🏠 **Remote:** {'Yes' if row.get('is_remote') else 'No'}")
                if has_ats:
                    _missing   = list(row.get("missing_skills") or [])
                    _matched_n = len(row.get("matched_skills") or [])
                    _total_n   = _matched_n + len(_missing)
                    m3.markdown(
                        f"🎯 **ATS Score:** {ats:.0f}%"
                        + (f" ({_matched_n}/{_total_n} keywords)" if _total_n else "")
                    )
                    # Why the score is what it is — single element per card
                    if _total_n == 0:
                        st.caption(
                            "🎯 **Why ATS 0%:** no skill keywords could be detected "
                            "for this job (no description was available to scan), "
                            "so keyword coverage can't be measured."
                        )
                    elif _missing:
                        _shown = ", ".join(f"`{s}`" for s in _missing[:10])
                        _extra = f" *+{len(_missing) - 10} more*" if len(_missing) > 10 else ""
                        st.caption(
                            f"🎯 **Why ATS {ats:.0f}%:** {len(_missing)} of {_total_n} "
                            f"job keywords are missing from your CV — {_shown}{_extra}"
                        )
                    else:
                        st.caption(
                            f"🎯 **Why ATS 100%:** your CV covers all {_total_n} "
                            "skill keywords detected in this job."
                        )
                _sal = row.get("salary_text")
                if isinstance(_sal, str) and _sal.strip():
                    m3.markdown(f"💰 **Salary:** {_sal}")

                url = str(row.get("job_url") or "")
                ac1, ac2, _ac3 = st.columns([2, 3, 5])
                # http(s) only — scraped hrefs could carry javascript:/data: schemes
                if url.lower().startswith(("http://", "https://")):
                    ac1.link_button("Apply Now ↗", url)

                # ── Tracker status (persisted to job_tracker.json) ───────────
                _cur = tracker.get(url, {}).get("status", "Not tracked")
                if _cur not in TRACK_STATUSES:
                    _cur = "Not tracked"
                _choice = ac2.selectbox(
                    "Track status", TRACK_STATUSES,
                    index=TRACK_STATUSES.index(_cur),
                    key=f"trk_{row.name}",
                    label_visibility="collapsed",
                    help="Track this job: Saved → Applied → Interview → Offer/Rejected",
                )
                if _choice != _cur and url:
                    if _choice == "Not tracked":
                        tracker.pop(url, None)
                    else:
                        _e = tracker.get(url, {})
                        _e.update({
                            "title":   str(row.get("title", "")),
                            "company": str(row.get("company", "")),
                            "site":    str(row.get("site", "")),
                            "status":  _choice,
                            "added":   _e.get("added") or datetime.now().strftime("%Y-%m-%d"),
                        })
                        tracker[url] = _e
                    _save_json(_TRACKER_PATH, tracker)

                desc = str(row.get("description", "") or "")
                if desc:
                    st.write(desc[:900] + ("…" if len(desc) > 900 else ""))

                # ── CV Tips ──────────────────────────────────────────────────
                # Works with or without a CV upload; no API key required.
                st.divider()
                st.markdown("##### 📝 CV Tips for this role")

                job_skills_val = row.get("job_skills")
                skills_in_job = (
                    sorted(job_skills_val)
                    if isinstance(job_skills_val, (list, set))
                    else sorted(extract_skills(desc))
                )

                if not skills_in_job:
                    st.caption("No recognisable skill keywords found in this job description.")
                else:
                    if st.session_state.cv_text:
                        # Split into already-present vs missing
                        already    = [s for s in skills_in_job if s in cv_skills]
                        missing_kw = [s for s in skills_in_job if s not in cv_skills]

                        # One markdown element per column instead of one per
                        # skill — Streamlit serializes every element on each
                        # rerun, so hundreds of tiny markdowns add real lag.
                        c_miss, c_have = st.columns(2)
                        if missing_kw:
                            c_miss.markdown(
                                "**Add to your CV** _(in the job, not in your CV)_\n\n"
                                + "\n".join(f"- `{s}`" for s in missing_kw)
                            )
                        else:
                            c_miss.success("Your CV already covers all detected skills.")

                        if already:
                            c_have.markdown(
                                "**Already in your CV** ✓\n\n"
                                + "\n".join(f"- ~~`{s}`~~" for s in already)
                            )
                    else:
                        # No CV — just show what the job requires so user can add manually
                        st.caption(
                            "Upload your CV in the sidebar to see which of these "
                            "you already have. For now, here's what this job asks for:"
                        )
                        cols = st.columns(3)
                        for i in range(3):
                            chunk = skills_in_job[i::3]
                            if chunk:
                                cols[i].markdown("\n".join(f"- `{s}`" for s in chunk))

                # ── AI actions (need CV + API key) ───────────────────────────
                if has_ai_key and st.session_state.cv_text:
                    _label  = ai_provider.split(" ")[0]
                    _title  = str(row.get("title", ""))
                    _co     = str(row.get("company", ""))
                    missing_for_ai = (
                        list(row.get("missing_skills") or [])
                        if has_scores
                        else [s for s in skills_in_job if s not in cv_skills]
                    )

                    b1, b2, b3, b4 = st.columns(4)

                    if b1.button("💡 CV suggestions", key=f"ai_{row.name}"):
                        with st.spinner(f"Asking {_label} …"):
                            result, ai_err = ai_suggestions(
                                st.session_state.cv_text, _title, desc,
                                missing_for_ai, api_key=api_key, provider=provider_key,
                            )
                        if result:
                            st.session_state[f"ai_res_{row.name}"] = result
                        else:
                            st.error(f"CV suggestions failed — {ai_err}")

                    if b2.button("✉️ Cover letter", key=f"cl_{row.name}"):
                        with st.spinner(f"Drafting cover letter with {_label} …"):
                            result, ai_err = ai_cover_letter(
                                st.session_state.cv_text, _title, _co, desc,
                                api_key=api_key, provider=provider_key,
                            )
                        if result:
                            st.session_state[f"cl_res_{row.name}"] = result
                        else:
                            st.error(f"Cover letter failed — {ai_err}")

                    if b3.button("📄 Tailored CV", key=f"cvd_{row.name}"):
                        with st.spinner(f"Rewriting CV with {_label} … (can take a minute)"):
                            result, ai_err = ai_tailored_cv(
                                st.session_state.cv_text, _title, desc,
                                missing_for_ai, api_key=api_key, provider=provider_key,
                            )
                        if result:
                            try:
                                st.session_state[f"cv_docx_{row.name}"] = _text_to_docx_bytes(result)
                            except Exception as e:
                                st.error(f"Could not build .docx — {e}")
                        else:
                            st.error(f"Tailored CV failed — {ai_err}")

                    if b4.button("🎤 Interview prep", key=f"ip_{row.name}"):
                        with st.spinner(f"Preparing questions with {_label} …"):
                            result, ai_err = ai_interview_prep(
                                st.session_state.cv_text, _title, desc,
                                missing_for_ai, api_key=api_key, provider=provider_key,
                            )
                        if result:
                            st.session_state[f"ip_res_{row.name}"] = result
                        else:
                            st.error(f"Interview prep failed — {ai_err}")

                    if st.session_state.get(f"ai_res_{row.name}"):
                        st.info(st.session_state[f"ai_res_{row.name}"])
                    if st.session_state.get(f"cl_res_{row.name}"):
                        st.text_area(
                            "Cover letter (copy from here)",
                            st.session_state[f"cl_res_{row.name}"],
                            height=260, key=f"cl_view_{row.name}",
                        )
                    if st.session_state.get(f"cv_docx_{row.name}"):
                        st.download_button(
                            "⬇️ Download tailored CV (.docx)",
                            data=st.session_state[f"cv_docx_{row.name}"],
                            file_name=f"CV_{_safe_filename(_co or _title)}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"cv_dl_{row.name}",
                        )
                    if st.session_state.get(f"ip_res_{row.name}"):
                        st.info(st.session_state[f"ip_res_{row.name}"])

    # ── Tab 2 — CV Analysis ──────────────────────────────────────────────────
    with tab_analysis:
        if not st.session_state.cv_text:
            st.info("Upload your CV in the sidebar to unlock match analysis.")
        elif not has_scores:
            st.info("Run a search to see match results.")
        else:
            pcts = df["match_score"].tolist()
            avg  = sum(pcts) / len(pcts) if pcts else 0

            has_ats_col = "ats_score" in df.columns

            m1, m2, m3, m4, m5 = st.columns(5)
            m1.metric("Total Jobs",    len(df))
            m2.metric("Avg Match",     f"{avg:.1f}%")
            m3.metric("Avg ATS",       f"{df['ats_score'].mean():.1f}%" if has_ats_col else "—")
            m4.metric("High  ≥65%",    int(sum(1 for p in pcts if p >= 65)))
            m5.metric("Medium 35–65%", int(sum(1 for p in pcts if 35 <= p < 65)))

            # Top 20 bar chart
            chart_df = (
                df[["title", "company", "match_score"]]
                .assign(label=df["title"].str[:28] + " @ " + df["company"].str[:18])
                .sort_values("match_score", ascending=False)
                .head(20)
            )
            st.bar_chart(chart_df.set_index("label")["match_score"])

            # Jobs-per-keyword breakdown (only when multi-keyword search)
            if "searched_keyword" in df.columns and df["searched_keyword"].nunique() > 1:
                st.subheader("Results by Job Title")
                kw_counts = df.groupby("searched_keyword").size().reset_index(name="count")
                st.bar_chart(kw_counts.set_index("searched_keyword")["count"])

            # Top 10 table
            st.subheader("Top 10 Matching Jobs")
            top_cols = ["title", "company", "location", "match_score", "missing_skills"]
            if has_ats_col:
                top_cols.insert(4, "ats_score")
            if "searched_keyword" in df.columns:
                top_cols = ["searched_keyword"] + top_cols
            top = df.nlargest(10, "match_score")[top_cols].copy()
            top["match_score"]   = top["match_score"].apply(lambda x: f"{x:.1f}%")
            if has_ats_col:
                top["ats_score"] = top["ats_score"].apply(lambda x: f"{x:.0f}%")
            top["missing_skills"] = top["missing_skills"].apply(
                lambda x: ", ".join(x[:5]) if isinstance(x, list) else ""
            )
            st.dataframe(top.reset_index(drop=True), use_container_width=True)

            st.divider()
            st.subheader("CV Preview")
            _cv = st.session_state.cv_text
            st.text_area(
                "CV text",
                _cv[:2000] + ("\n…" if len(_cv) > 2000 else ""),
                height=200,
                disabled=True,
                label_visibility="collapsed",
            )

    # ── Tab 3 — Tracker ──────────────────────────────────────────────────────
    with tab_tracker:
        if not tracker:
            st.info(
                "No tracked jobs yet. Open any job card and set its status "
                "(Saved / Applied / Interview / Offer / Rejected) next to the Apply button."
            )
        else:
            _by_status: dict = {}
            for _u, _e in tracker.items():
                _by_status.setdefault(_e.get("status", "Saved"), []).append((_u, _e))

            mt = st.columns(5)
            for _i, _s in enumerate(["Saved", "Applied", "Interview", "Offer", "Rejected"]):
                mt[_i].metric(_s, len(_by_status.get(_s, [])))
            st.divider()

            for _u, _e in sorted(tracker.items(), key=lambda kv: kv[1].get("added", ""), reverse=True):
                tc1, tc2, tc3, tc4 = st.columns([5, 2, 2, 1])
                tc1.markdown(
                    f"**[{_e.get('title', 'Job')}]({_u})** — {_e.get('company', '—')}  "
                    f"· {_e.get('site', '')} · added {_e.get('added', '')}"
                )
                _hk = abs(hash(_u))
                _opts = TRACK_STATUSES[1:]
                _cur_t = _e.get("status", "Saved")
                _new_t = tc2.selectbox(
                    "Status", _opts,
                    index=_opts.index(_cur_t) if _cur_t in _opts else 0,
                    key=f"trkt_{_hk}", label_visibility="collapsed",
                )
                if _new_t != _cur_t:
                    tracker[_u]["status"] = _new_t
                    _save_json(_TRACKER_PATH, tracker)
                    st.rerun()
                if tc4.button("🗑", key=f"trkdel_{_hk}", help="Remove from tracker"):
                    tracker.pop(_u, None)
                    _save_json(_TRACKER_PATH, tracker)
                    st.rerun()

    # ── Tab 4 — Download ─────────────────────────────────────────────────────
    with tab_download:
        st.subheader("Download Results")

        if st.button("Generate Formatted Excel Report  📊", type="primary"):
            # Build match_scores list expected by excel_exporter
            if has_scores:
                scores_for_export = df[
                    ["match_score", "matched_skills", "missing_skills"]
                ].rename(columns={"match_score": "tfidf_score"}).to_dict("records")
            else:
                scores_for_export = None

            with st.spinner("Building Excel …"):
                with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as tmp:
                    tmp_path = tmp.name
                try:
                    export_to_excel(df, scores_for_export, output_path=tmp_path)
                    with open(tmp_path, "rb") as fh:
                        st.session_state.excel_bytes = fh.read()
                finally:
                    if os.path.exists(tmp_path):
                        os.unlink(tmp_path)

        if st.session_state.excel_bytes:
            ts = datetime.now().strftime("%Y-%m-%d_%H%M")
            st.download_button(
                "⬇️  Download Excel  (.xlsx)",
                data=st.session_state.excel_bytes,
                file_name=f"jobs_{ts}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True,
            )

        st.divider()
        st.subheader("Raw CSV")
        _kw_str  = "_".join(keywords[:3]) if keywords else "jobs"
        safe_kw  = _safe_filename(_kw_str)
        safe_loc = _safe_filename(location or "")
        csv_data = _csv_safe(df.drop(
            columns=["matched_skills", "missing_skills", "job_skills"], errors="ignore"
        )).to_csv(index=False).encode("utf-8")
        st.download_button(
            "⬇️  Download CSV",
            data=csv_data,
            file_name=f"{safe_kw}_{safe_loc}.csv",
            mime="text/csv",
            use_container_width=True,
        )

# ── Empty / welcome state ─────────────────────────────────────────────────────
else:
    st.title("💼 Job Finder AI")
    st.markdown(
        "Search **LinkedIn, Indeed, Remotive, Jobicy, The Muse, Arbeitnow, Naukri, "
        "Seek, JobStreet, Bayt, Reed, Rozee** and more — no login, no API key.  \n"
        "Job boards are **auto-suggested** based on the country you pick."
    )

    with st.expander("📖 Quick start", expanded=True):
        st.markdown(
            """
| Step | What to do |
|------|-----------|
| 1 | **Select one or more job titles** from the list, or type custom ones (one per line) |
| 2 | **Pick a country** — job boards update automatically |
| 3 | Optionally narrow to a **city** |
| 4 | Adjust job boards, date window, and max jobs |
| 5 | *(Optional)* Upload your CV for per-job match scores |
| 6 | *(Optional)* Pick an AI provider and enter your key for CV suggestions |
| 7 | Click **Search Jobs 🚀** — all titles and boards searched in parallel |
| 8 | Filter by title, site, type, remote, or match score |
| 9 | Download formatted **Excel** or **CSV** |

**Platform coverage by country:**
| Country | Boards |
|---------|--------|
| Remote (Worldwide) | LinkedIn, RemoteOK, Remotive, Jobicy |
| United States / Canada | LinkedIn, Indeed, The Muse |
| India | LinkedIn, Indeed, Naukri |
| Australia / NZ | LinkedIn, Seek |
| UK | LinkedIn, Indeed, Reed |
| Germany / Austria | LinkedIn, Indeed, Arbeitnow |
| UAE / Saudi / Gulf | LinkedIn, Bayt |
| Malaysia / Philippines / Singapore / Indonesia | LinkedIn, JobStreet |
| Pakistan | LinkedIn, Rozee.pk |
| All others | LinkedIn, Indeed |
            """
        )
    st.info("Select options in the sidebar and click **Search Jobs 🚀** to begin.")

    _tracker_home = _load_json(_TRACKER_PATH, {})
    if _tracker_home:
        with st.expander(f"📌 Your tracked jobs ({len(_tracker_home)})", expanded=False):
            for _u, _e in sorted(_tracker_home.items(),
                                 key=lambda kv: kv[1].get("added", ""), reverse=True):
                st.markdown(
                    f"- **[{_e.get('title', 'Job')}]({_u})** — {_e.get('company', '—')} "
                    f"· `{_e.get('status', 'Saved')}` · added {_e.get('added', '')}"
                )
