"""
CV analysis module.

Reads DOCX/PDF CVs, scores them against job descriptions using TF-IDF cosine
similarity, extracts matching/missing keywords, and optionally queries Claude
for tailored improvement suggestions.
"""

import re
import os
import logging
from typing import Optional

import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

logger = logging.getLogger(__name__)

# ──────────────────────────────────────────────────────────────────────────────
# CV READING
# ──────────────────────────────────────────────────────────────────────────────

def read_cv(filepath: str) -> str:
    """
    Extract all text from a CV file (.docx or .pdf).

    Returns the raw text string, or an empty string on failure.
    """
    ext = filepath.rsplit(".", 1)[-1].lower()
    if ext == "docx":
        return _read_docx(filepath)
    elif ext == "pdf":
        return _read_pdf(filepath)
    else:
        raise ValueError(f"Unsupported format: .{ext}  — use .docx or .pdf")


def _read_docx(filepath: str) -> str:
    try:
        from docx import Document
        doc = Document(filepath)
        parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        # Also pull text from tables (skills tables, etc.)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    t = cell.text.strip()
                    if t:
                        parts.append(t)
        return "\n".join(parts)
    except Exception as e:
        logger.error(f"DOCX read error: {e}")
        return ""


def _read_pdf(filepath: str) -> str:
    # pdfplumber preserves layout better than PyPDF2
    try:
        import pdfplumber
        parts = []
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    parts.append(text)
        return "\n".join(parts)
    except ImportError:
        logger.warning("pdfplumber not installed — falling back to PyPDF2")
    except Exception as e:
        logger.error(f"pdfplumber error: {e}")

    # Fallback
    try:
        import PyPDF2
        with open(filepath, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            return "\n".join(page.extract_text() or "" for page in reader.pages)
    except ImportError:
        logger.error("Neither pdfplumber nor PyPDF2 is installed")
    except Exception as e:
        logger.error(f"PyPDF2 error: {e}")

    return ""


# ──────────────────────────────────────────────────────────────────────────────
# SIMILARITY SCORING
# ──────────────────────────────────────────────────────────────────────────────

_ST_MODEL = None
_ST_MODEL_NAME = ""
_ST_UNAVAILABLE = False   # cache the failed import — retrying it scans sys.path every call

# Best small free model on MTEB that is also SAFE to load: standard BERT
# architecture, safetensors weights, MIT license, and crucially loads WITHOUT
# trust_remote_code (unlike nomic/jina, which execute repo code on load).
# ~62 MTEB avg vs ~56 for all-MiniLM-L6-v2 at a similar size (33M params).
_DEFAULT_EMBED_MODEL  = "BAAI/bge-small-en-v1.5"
_FALLBACK_EMBED_MODEL = "all-MiniLM-L6-v2"   # already cached locally


def _get_st_model():
    """Load the sentence-transformers model once per process.
    Model can be overridden with the JOBFINDER_EMBED_MODEL env var."""
    global _ST_MODEL, _ST_MODEL_NAME, _ST_UNAVAILABLE
    if _ST_UNAVAILABLE:
        raise ImportError("sentence-transformers unavailable (cached)")
    if _ST_MODEL is None:
        try:
            from sentence_transformers import SentenceTransformer
        except ImportError:
            _ST_UNAVAILABLE = True
            raise
        name = os.getenv("JOBFINDER_EMBED_MODEL", _DEFAULT_EMBED_MODEL)
        try:
            # trust_remote_code stays at its default (False) — never load a
            # model that requires executing code from the model repo.
            _ST_MODEL = SentenceTransformer(name)
            _ST_MODEL_NAME = name
        except Exception as e:
            logger.warning(
                f"Could not load embedding model '{name}' ({e}) — "
                f"falling back to {_FALLBACK_EMBED_MODEL}"
            )
            _ST_MODEL = SentenceTransformer(_FALLBACK_EMBED_MODEL)
            _ST_MODEL_NAME = _FALLBACK_EMBED_MODEL
    return _ST_MODEL


def _api_embed_scores(cv_text: str, descriptions: list[str],
                      api_key: str, provider: str) -> list[float]:
    """Score with a provider's embeddings API (higher quality than any small
    local model). Raises on any failure — caller falls back to local."""
    texts = [cv_text[:8000]] + [str(d)[:4000] for d in descriptions]

    if provider == "openai":
        from openai import OpenAI
        client = OpenAI(api_key=api_key)
        resp = client.embeddings.create(model="text-embedding-3-small", input=texts)
        embs = np.array([d.embedding for d in resp.data])
    elif provider == "gemini":
        import google.generativeai as genai
        genai.configure(api_key=api_key)
        r = genai.embed_content(
            model="models/text-embedding-004",
            content=texts,
            task_type="semantic_similarity",
        )
        embs = np.array(r["embedding"])
    else:
        # Anthropic has no embeddings endpoint — a Claude key can't embed
        raise ValueError(f"provider '{provider}' has no embeddings API")

    sims = cosine_similarity(embs[:1], embs[1:])[0]
    return [round(float(s) * 100, 1) for s in sims]


def score_cv_against_jobs(cv_text: str, descriptions: list[str],
                          api_key: str = None, provider: str = None) -> list[float]:
    """
    Compute a 0-100 similarity score between the CV and each job description.

    Priority: provider embeddings API (when an OpenAI/Gemini key is given)
    → local sentence-transformers (bge-small-en-v1.5) → TF-IDF.
    """
    if not cv_text or not descriptions:
        return [0.0] * len(descriptions)

    # 1) API embeddings when an embedding-capable key is available
    if api_key and provider in ("openai", "gemini"):
        try:
            return _api_embed_scores(cv_text, descriptions, api_key, provider)
        except Exception as e:
            logger.warning(f"{provider} embeddings failed ({e}) — using local model")

    # 2) Local sentence-transformers (semantic understanding)
    try:
        model = _get_st_model()

        # The model truncates input around 256 tokens, so embed the CV in
        # chunks and mean-pool — otherwise everything past page 1 is ignored.
        chunks = [cv_text[i:i + 1000] for i in range(0, min(len(cv_text), 5000), 1000)]
        cv_emb = model.encode(chunks, convert_to_numpy=True).mean(axis=0, keepdims=True)
        job_embs = model.encode(
            [str(d)[:1024] for d in descriptions], convert_to_numpy=True
        )

        sims = cosine_similarity(cv_emb, job_embs)[0]
        if "bge" in _ST_MODEL_NAME.lower():
            # bge cosine similarity is compressed into roughly [0.5, 0.95]
            # (unrelated text still scores ~0.55) — rescale so the UI's
            # 35%/65% badge thresholds keep their meaning.
            sims = np.clip((sims - 0.5) / 0.45, 0.0, 1.0)
        return [round(float(s) * 100, 1) for s in sims]

    except ImportError:
        logger.info("sentence-transformers not installed — using TF-IDF")
    except Exception as e:
        logger.warning(f"sentence-transformers failed ({e}), falling back to TF-IDF")

    return _tfidf_score(cv_text, descriptions)


def _tfidf_score(cv_text: str, descriptions: list[str]) -> list[float]:
    """TF-IDF cosine similarity fallback."""
    texts = [cv_text] + [str(d) for d in descriptions]
    try:
        vec = TfidfVectorizer(
            stop_words="english",
            ngram_range=(1, 2),
            max_features=8000,
            sublinear_tf=True,
        )
        mat = vec.fit_transform(texts)
        sims = cosine_similarity(mat[0], mat[1:])[0]
        return [round(float(s) * 100, 1) for s in sims]
    except Exception as e:
        logger.error(f"TF-IDF error: {e}")
        return [0.0] * len(descriptions)


# ──────────────────────────────────────────────────────────────────────────────
# KEYWORD GAP ANALYSIS
# ──────────────────────────────────────────────────────────────────────────────

# Comprehensive skill keyword list
_SKILLS = [
    # Languages
    "python", "java", "javascript", "typescript", "c++", "c#", "go", "rust",
    "ruby", "php", "swift", "kotlin", "scala", "r", "matlab", "bash", "sql",
    "html", "css", "dart", "perl", "groovy", "vba",
    # Frameworks / Libraries
    "react", "angular", "vue", "next.js", "nuxt", "svelte",
    "django", "flask", "fastapi", "spring", "spring boot",
    "node.js", "express", "laravel", "rails", "asp.net",
    "tensorflow", "pytorch", "keras", "scikit-learn", "pandas", "numpy",
    "hugging face", "langchain", "llm", "openai",
    # Cloud & DevOps
    "aws", "azure", "gcp", "google cloud", "docker", "kubernetes", "terraform",
    "ansible", "jenkins", "github actions", "gitlab ci", "circleci",
    "ci/cd", "devops", "devsecops", "helm", "prometheus", "grafana", "datadog",
    # Databases
    "postgresql", "mysql", "mongodb", "sqlite", "redis", "elasticsearch",
    "cassandra", "dynamodb", "firebase", "bigquery", "snowflake",
    # Data / AI
    "machine learning", "deep learning", "nlp", "computer vision",
    "data science", "data engineering", "data analysis", "etl",
    "tableau", "power bi", "looker", "dbt", "spark", "hadoop", "airflow",
    "mlflow", "feature engineering", "a/b testing", "statistics",
    # Tools
    "git", "github", "gitlab", "jira", "confluence", "figma",
    "postman", "swagger", "linux", "bash scripting",
    # Methodologies
    "agile", "scrum", "kanban", "tdd", "bdd", "rest", "graphql",
    "microservices", "event driven", "grpc",
    # Business
    "project management", "product management", "leadership",
    "communication", "excel", "powerpoint", "salesforce", "crm",
]

def _compile_skill(s: str) -> re.Pattern:
    # \b fails next to "+" / "#" (non-word chars), so "c++" and "c#" need
    # lookarounds instead of plain word boundaries.
    # No IGNORECASE: extract_skills lowercases the text once, and these
    # patterns run ~180× per description — case-sensitive matching is faster.
    esc = re.escape(s)
    return re.compile(rf"(?<![\w+#]){esc}(?![\w+#])")


# Pre-compile patterns once for speed
_SKILL_PATTERNS = {s: _compile_skill(s) for s in _SKILLS}

# Ambiguous one-word skills need context or they match plain English.
_SKILL_PATTERNS["r"] = re.compile(r"\br\b(?!\s*&)")  # skip "R&D"
_SKILL_PATTERNS["go"] = re.compile(
    r"\bgolang\b|\bgo\b(?=\s*(?:[,/;)|]|$|\s(?:programming|developer|language|lang|engineer)))",
    re.MULTILINE,
)


def extract_skills(text: str) -> set:
    """Return the set of known skills mentioned in the text."""
    tl = text.lower()
    return {s for s, pat in _SKILL_PATTERNS.items() if pat.search(tl)}


def keyword_gap(cv_text: str, job_description: str, cv_skills: set = None) -> dict:
    """
    Return which skills appear in the job description but not in the CV
    (and which ones do appear).

    Pass a precomputed cv_skills set (from extract_skills) when scoring many
    jobs against the same CV — avoids re-scanning the CV once per job.
    """
    in_job = extract_skills(job_description)
    if cv_skills is None:
        cv_skills = extract_skills(cv_text)
    in_cv = in_job & cv_skills
    missing = in_job - in_cv

    keyword_score = round(len(in_cv) / len(in_job) * 100, 1) if in_job else 0.0

    return {
        "keyword_score": keyword_score,
        "matched_skills": sorted(in_cv),
        "missing_skills": sorted(missing),
        "total_job_skills": len(in_job),
    }


# ──────────────────────────────────────────────────────────────────────────────
# AI SUGGESTIONS  (Claude · ChatGPT · Gemini — all optional)
# ──────────────────────────────────────────────────────────────────────────────

_PROVIDER_ENV_VARS: dict[str, str] = {
    "claude": "ANTHROPIC_API_KEY",
    "openai": "OPENAI_API_KEY",
    "gemini": "GOOGLE_API_KEY",
}


def _resolve_key(api_key: Optional[str], provider: str) -> tuple[Optional[str], str]:
    env_var = _PROVIDER_ENV_VARS.get(provider, "ANTHROPIC_API_KEY")
    return api_key or os.getenv(env_var), env_var


def _complete(prompt: str, key: str, provider: str,
              max_tokens: int = 4000) -> tuple[Optional[str], Optional[str]]:
    """Run one completion against the chosen provider.
    Returns (text, None) on success or (None, error_message) on failure."""
    if provider == "openai":
        return _openai_complete(prompt, key, max_tokens)
    elif provider == "gemini":
        return _gemini_complete(prompt, key)
    return _claude_complete(prompt, key, max_tokens)


def ai_suggestions(
    cv_text: str,
    job_title: str,
    job_description: str,
    missing_skills: list[str],
    api_key: str = None,
    provider: str = "claude",
) -> tuple[Optional[str], Optional[str]]:
    """
    Get AI-powered CV improvement suggestions.

    provider: "claude" | "openai" | "gemini"
    Returns (suggestion_text, None) on success or (None, error_message) on
    failure — the error is shown in the UI so a bad key / missing package
    doesn't fail silently.
    """
    key, env_var = _resolve_key(api_key, provider)
    if not key:
        return None, f"No API key — enter one in the sidebar or set {env_var}."

    prompt = _build_ai_prompt(cv_text, job_title, job_description, missing_skills)
    return _complete(prompt, key, provider, max_tokens=4000)


def ai_cover_letter(
    cv_text: str,
    job_title: str,
    company: str,
    job_description: str,
    api_key: str = None,
    provider: str = "claude",
) -> tuple[Optional[str], Optional[str]]:
    """Draft a tailored cover letter for one job. Same return contract as
    ai_suggestions."""
    key, env_var = _resolve_key(api_key, provider)
    if not key:
        return None, f"No API key — enter one in the sidebar or set {env_var}."

    prompt = f"""Write a tailored cover letter for this job application.

Role: {job_title}
Company: {company or "the company"}

Job Description (excerpt):
{job_description[:1800]}

My CV (excerpt):
{cv_text[:2500]}

Requirements:
- 3 short paragraphs, under 250 words total — hiring managers skim
- Open with a specific hook connecting my strongest relevant experience to this role
- Reference 2-3 concrete achievements from my CV that map to the job's needs
- Mention ONLY experience that is actually in my CV — never invent anything
- Professional but human tone, no clichés like "I am writing to express"
- End with a confident, brief closing
- Output the letter only — no preamble, no placeholders other than [Hiring Manager] if no name is known"""
    return _complete(prompt, key, provider, max_tokens=3000)


def ai_tailored_cv(
    cv_text: str,
    job_title: str,
    job_description: str,
    missing_skills: list[str],
    api_key: str = None,
    provider: str = "claude",
) -> tuple[Optional[str], Optional[str]]:
    """Rewrite the full CV text targeted at one job (for .docx export).
    Same return contract as ai_suggestions."""
    key, env_var = _resolve_key(api_key, provider)
    if not key:
        return None, f"No API key — enter one in the sidebar or set {env_var}."

    missing_str = ", ".join(missing_skills[:12]) if missing_skills else "none identified"
    prompt = f"""Rewrite my CV tailored for this specific role. Output the complete rewritten CV as plain text.

Target role: {job_title}

Job Description (excerpt):
{job_description[:1800]}

My current CV:
{cv_text[:5000]}

Keywords from the job missing from my CV: {missing_str}

Rules:
- Keep ALL factual content truthful — never invent skills, employers, dates, or experience
- Reorder and reword so the most relevant experience for THIS role leads
- Weave the missing keywords in ONLY where my real experience genuinely supports them
- Quantify achievements where the original CV gives numbers
- Use clear section headings (SUMMARY, EXPERIENCE, SKILLS, EDUCATION, ...) in UPPERCASE on their own line
- Use "- " bullets for achievements
- Output the CV text only — no commentary before or after"""
    return _complete(prompt, key, provider, max_tokens=8000)


def ai_interview_prep(
    cv_text: str,
    job_title: str,
    job_description: str,
    missing_skills: list[str],
    api_key: str = None,
    provider: str = "claude",
) -> tuple[Optional[str], Optional[str]]:
    """Generate likely interview questions + prep pointers for one job.
    Same return contract as ai_suggestions."""
    key, env_var = _resolve_key(api_key, provider)
    if not key:
        return None, f"No API key — enter one in the sidebar or set {env_var}."

    missing_str = ", ".join(missing_skills[:12]) if missing_skills else "none identified"
    prompt = f"""Prepare me for an interview for this role.

Role: {job_title}

Job Description (excerpt):
{job_description[:1800]}

My CV (excerpt):
{cv_text[:2200]}

Skills the job wants that my CV doesn't show: {missing_str}

Provide:
1. The 6 most likely technical/role-specific questions for THIS job, each with a 1-2 sentence answer strategy based on MY actual experience
2. The 2 hardest questions I'll face about my gaps ({missing_str}) and honest ways to handle them
3. Two strong questions I should ask the interviewer, specific to this role

Be direct and specific to this job and my background. No generic interview advice."""
    return _complete(prompt, key, provider, max_tokens=4000)


def _build_ai_prompt(cv_text, job_title, job_description, missing_skills) -> str:
    missing_str = ", ".join(missing_skills[:12]) if missing_skills else "none identified"
    return f"""You are an expert CV/resume optimizer. Give specific, actionable advice.

Role: {job_title}

Job Description (excerpt):
{job_description[:1800]}

My CV (excerpt):
{cv_text[:2200]}

Skills in the job description missing from my CV: {missing_str}

Provide:
1. Top 3 concrete changes to make this CV stronger for this specific role
2. Keywords / phrases to add naturally (with examples of where to place them)
3. Rewrite 2 of my actual CV bullet points for this role — quoted before/after,
   quantified where possible, weaving in the missing keywords naturally
4. One sentence on overall fit: strong / moderate / weak and why

Be direct, specific, and concise. Base rewrites only on experience actually
present in my CV — never invent skills or experience I don't have."""


def _claude_complete(prompt: str, api_key: str, max_tokens: int = 4000) -> tuple[Optional[str], Optional[str]]:
    try:
        import anthropic
    except ImportError:
        logger.warning("anthropic not installed — run: pip install anthropic")
        return None, "Python package 'anthropic' is not installed — run: pip install anthropic"
    try:
        client = anthropic.Anthropic(api_key=api_key)
        msg = client.messages.create(
            model="claude-opus-4-8",
            max_tokens=max_tokens,
            thinking={"type": "adaptive"},
            messages=[{"role": "user", "content": prompt}],
        )
        # With thinking enabled, content[0] may be a thinking block —
        # pick the text block instead of indexing blindly.
        text = next((b.text for b in msg.content if b.type == "text"), None)
        if text:
            return text, None
        return None, f"Claude returned no text (stop_reason: {msg.stop_reason})."
    except anthropic.AuthenticationError:
        return None, "Invalid Anthropic API key — check the key in the sidebar."
    except anthropic.RateLimitError:
        return None, "Anthropic rate limit hit — wait a minute and retry."
    except Exception as e:
        logger.error(f"Claude API error: {e}")
        return None, f"Claude API error: {e}"


def _openai_complete(prompt: str, api_key: str, max_tokens: int = 2000) -> tuple[Optional[str], Optional[str]]:
    try:
        from openai import OpenAI
    except ImportError:
        logger.warning("openai not installed — run: pip install openai")
        return None, "Python package 'openai' is not installed — run: pip install openai"
    try:
        client = OpenAI(api_key=api_key)
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            max_tokens=max_tokens,
            messages=[{"role": "user", "content": prompt}],
        )
        return response.choices[0].message.content, None
    except Exception as e:
        logger.error(f"OpenAI API error: {e}")
        return None, f"OpenAI API error: {e}"


def _gemini_complete(prompt: str, api_key: str) -> tuple[Optional[str], Optional[str]]:
    try:
        import google.generativeai as genai
    except ImportError:
        logger.warning("google-generativeai not installed — run: pip install google-generativeai")
        return None, "Python package 'google-generativeai' is not installed — run: pip install google-generativeai"
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel("gemini-2.5-flash")
        response = model.generate_content(prompt)
        return response.text, None
    except Exception as e:
        logger.error(f"Gemini API error: {e}")
        return None, f"Gemini API error: {e}"
